from typing import Dict, Any, Optional, List
from datetime import datetime
import json
from pathlib import Path
from jinja2 import Template
import markdown2
import logging
import base64
from io import BytesIO

logger = logging.getLogger(__name__)

# Optional imports with fallbacks
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab not available, PDF generation will be limited")

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available, DOCX generation will be limited")

class ReportService:
    """Service for generating various report formats"""
    
    def __init__(self):
        self.reports_dir = Path("reports")
        self.reports_dir.mkdir(exist_ok=True)
        self.templates_dir = Path("templates")
        self.templates_dir.mkdir(exist_ok=True)
    
    async def generate_pdf(
        self,
        report_id: str,
        content: Dict[str, Any],
        include_charts: bool = True
    ) -> str:
        """Generate PDF report using ReportLab or fallback to HTML"""
        
        try:
            output_path = self.reports_dir / f"{report_id}.pdf"
            
            if REPORTLAB_AVAILABLE:
                # Use ReportLab for PDF generation
                return await self._generate_pdf_reportlab(output_path, content, include_charts)
            else:
                # Fallback to HTML format with PDF extension
                logger.warning("ReportLab not available, generating HTML-based report")
                html_content = await self._create_html_content(content, include_charts)
                
                # Save as HTML but with .pdf extension for compatibility
                html_path = self.reports_dir / f"{report_id}.html"
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                
                return str(html_path)
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            # Fallback to simple text report
            return await self._generate_text_fallback(report_id, content)
    
    async def generate_html(
        self,
        report_id: str,
        content: Dict[str, Any],
        include_charts: bool = True
    ) -> str:
        """Generate HTML report"""
        
        html_content = await self._create_html_content(content, include_charts)
        output_path = self.reports_dir / f"{report_id}.html"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return str(output_path)
    
    async def generate_markdown(
        self,
        report_id: str,
        content: Dict[str, Any]
    ) -> str:
        """Generate Markdown report"""
        
        md_content = self._create_markdown_content(content)
        output_path = self.reports_dir / f"{report_id}.md"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return str(output_path)
    
    async def generate_docx(
        self,
        report_id: str,
        content: Dict[str, Any],
        include_charts: bool = True
    ) -> str:
        """Generate Word document report"""
        
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, falling back to markdown")
            return await self.generate_markdown(report_id, content)
        
        try:
            doc = Document()
            
            # Add title
            doc.add_heading(content.get('title', 'Analysis Report'), 0)
            
            # Add metadata
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add executive summary
            if 'executive_summary' in content:
                doc.add_heading('Executive Summary', level=1)
                doc.add_paragraph(content['executive_summary'])
            
            # Add sections
            for section in content.get('sections', []):
                doc.add_heading(section['title'], level=1)
                doc.add_paragraph(section['content'])
                
                # Add data table if present
                if 'data' in section and isinstance(section['data'], list) and len(section['data']) > 0:
                    self._add_table_to_docx(doc, section['data'])
            
            output_path = self.reports_dir / f"{report_id}.docx"
            doc.save(output_path)
            
            return str(output_path)
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            return await self._generate_text_fallback(report_id, content)
    
    async def _create_html_content(
        self,
        content: Dict[str, Any],
        include_charts: bool
    ) -> str:
        """Create HTML content for report"""
        
        template = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{{ title }}</title>
            <style>
                body { font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; }
                h1 { color: #0078D4; border-bottom: 2px solid #0078D4; padding-bottom: 10px; }
                h2 { color: #333; margin-top: 30px; }
                .metadata { color: #666; font-size: 0.9em; }
                .summary { background: #f5f5f5; padding: 20px; border-left: 4px solid #0078D4; margin: 20px 0; }
                .section { margin: 30px 0; }
                table { border-collapse: collapse; width: 100%; margin: 20px 0; }
                th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }
                th { background: #0078D4; color: white; }
                .chart { margin: 20px 0; text-align: center; }
            </style>
        </head>
        <body>
            <h1>{{ title }}</h1>
            <p class="metadata">Generated: {{ generated_date }}</p>
            
            {% if executive_summary %}
            <div class="summary">
                <h2>Executive Summary</h2>
                <p>{{ executive_summary }}</p>
            </div>
            {% endif %}
            
            {% for section in sections %}
            <div class="section">
                <h2>{{ section.title }}</h2>
                <p>{{ section.content }}</p>
                
                {% if section.data %}
                <table>
                    <thead>
                        <tr>
                            {% for col in section.data[0].keys() %}
                            <th>{{ col }}</th>
                            {% endfor %}
                        </tr>
                    </thead>
                    <tbody>
                        {% for row in section.data %}
                        <tr>
                            {% for value in row.values() %}
                            <td>{{ value }}</td>
                            {% endfor %}
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
                {% endif %}
                
                {% if include_charts and section.chart %}
                <div class="chart">
                    <!-- Chart placeholder -->
                </div>
                {% endif %}
            </div>
            {% endfor %}
        </body>
        </html>
        """
        
        tmpl = Template(template)
        
        return tmpl.render(
            title=content.get('title', 'Report'),
            generated_date=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            executive_summary=content.get('executive_summary', ''),
            sections=content.get('sections', []),
            include_charts=include_charts
        )
    
    def _create_markdown_content(self, content: Dict[str, Any]) -> str:
        """Create Markdown content for report"""
        
        md_lines = []
        
        # Title
        md_lines.append(f"# {content.get('title', 'Report')}")
        md_lines.append(f"\n*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
        
        # Executive Summary
        if 'executive_summary' in content:
            md_lines.append("## Executive Summary\n")
            md_lines.append(content['executive_summary'])
            md_lines.append("\n")
        
        # Sections
        for section in content.get('sections', []):
            md_lines.append(f"## {section['title']}\n")
            md_lines.append(section['content'])
            md_lines.append("\n")
            
            # Add data table if present
            if 'data' in section and section['data']:
                # Create markdown table
                headers = list(section['data'][0].keys())
                md_lines.append("| " + " | ".join(headers) + " |")
                md_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                
                for row in section['data']:
                    md_lines.append("| " + " | ".join(str(v) for v in row.values()) + " |")
                
                md_lines.append("\n")
        
        return "\n".join(md_lines)
    
    async def _generate_pdf_reportlab(self, output_path: Path, content: Dict[str, Any], include_charts: bool) -> str:
        """Generate PDF using ReportLab"""
        
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#0078D4'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph(content.get('title', 'Analysis Report'), title_style))
        story.append(Spacer(1, 12))
        
        # Metadata
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Executive Summary
        if 'executive_summary' in content:
            story.append(Paragraph('Executive Summary', styles['Heading2']))
            story.append(Paragraph(content['executive_summary'], styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Sections
        for section in content.get('sections', []):
            story.append(Paragraph(section['title'], styles['Heading2']))
            story.append(Paragraph(section['content'], styles['Normal']))
            
            # Add table if present
            if 'data' in section and isinstance(section['data'], list) and len(section['data']) > 0:
                table_data = self._prepare_table_data(section['data'])
                if table_data:
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 14),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(t)
            
            story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        return str(output_path)
    
    def _prepare_table_data(self, data: List[Dict]) -> List[List]:
        """Prepare data for ReportLab table"""
        if not data:
            return []
        
        # Get headers
        headers = list(data[0].keys())
        table_data = [headers]
        
        # Add rows (limit to 50 rows for PDF)
        for row in data[:50]:
            table_data.append([str(row.get(h, '')) for h in headers])
        
        return table_data
    
    def _add_table_to_docx(self, doc, data: List[Dict]):
        """Add a table to DOCX document"""
        if not data or not isinstance(data, list):
            return
        
        # Get headers
        headers = list(data[0].keys())
        
        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = str(header)
        
        # Add data rows (limit to 50)
        for row_data in data[:50]:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                row_cells[i].text = str(row_data.get(header, ''))
    
    async def _generate_text_fallback(self, report_id: str, content: Dict[str, Any]) -> str:
        """Generate simple text report as fallback"""
        
        lines = []
        lines.append("=" * 60)
        lines.append(content.get('title', 'Analysis Report').center(60))
        lines.append("=" * 60)
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("")
        
        if 'executive_summary' in content:
            lines.append("EXECUTIVE SUMMARY")
            lines.append("-" * 40)
            lines.append(content['executive_summary'])
            lines.append("")
        
        for section in content.get('sections', []):
            lines.append(section['title'].upper())
            lines.append("-" * 40)
            lines.append(section['content'])
            lines.append("")
        
        output_path = self.reports_dir / f"{report_id}.txt"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
        
        return str(output_path)
